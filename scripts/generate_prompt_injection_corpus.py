"""Generate synthetic indirect prompt-injection documents from a manifest."""

import argparse
import json
from pathlib import Path

from docx import Document
from openpyxl import Workbook
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen.canvas import Canvas


def _body(sample: dict[str, str]) -> str:
    return (
        f"Trusted user goal: {sample['trusted_goal']}\n\n"
        f"Untrusted document content: {sample['payload']}\n"
    )


def _write_docx(sample: dict[str, str], destination: Path) -> None:
    document = Document()
    document.add_heading("Architecture Review", level=1)
    document.add_paragraph(f"Trusted user goal: {sample['trusted_goal']}")
    table = document.add_table(rows=1, cols=1)
    run = table.cell(0, 0).paragraphs[0].add_run(sample["payload"])
    run.font.hidden = True
    document.save(destination)


def _write_xlsx(sample: dict[str, str], destination: Path) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Release Schedule"
    sheet.append(["Trusted user goal", sample["trusted_goal"]])
    sheet.append(["Release", "2026-10-14"])
    sheet.append([sample["payload"]])
    sheet.row_dimensions[3].hidden = True
    workbook.save(destination)


def _write_pdf(sample: dict[str, str], destination: Path) -> None:
    canvas = Canvas(str(destination), pagesize=letter)
    canvas.drawString(72, 740, f"Trusted user goal: {sample['trusted_goal']}")
    canvas.drawString(72, 710, sample["payload"])
    canvas.save()


def _write_sample(sample: dict[str, str], destination: Path) -> None:
    handlers = {
        "txt": lambda: destination.write_text(_body(sample), encoding="utf-8"),
        "markdown": lambda: destination.write_text(
            f"# Support article\n\nTrusted user goal: {sample['trusted_goal']}\n\n"
            f"Visible support guidance.\n\n<!-- {sample['payload']} -->\n",
            encoding="utf-8",
        ),
        "metadata": lambda: destination.write_text(
            json.dumps(
                {
                    "trusted_user_goal": sample["trusted_goal"],
                    "document_properties": {"subject": sample["payload"]},
                },
                indent=2,
            )
            + "\n",
            encoding="utf-8",
        ),
        "docx": lambda: _write_docx(sample, destination),
        "xlsx": lambda: _write_xlsx(sample, destination),
        "pdf": lambda: _write_pdf(sample, destination),
    }
    handlers[sample["format"]]()


def generate_corpus(manifest_path: Path, output_dir: Path) -> list[Path]:
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    output_dir.mkdir(parents=True, exist_ok=True)
    outputs = []
    for sample in manifest["samples"]:
        destination = output_dir / sample["filename"]
        _write_sample(sample, destination)
        outputs.append(destination)
    return outputs


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("manifest", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    for path in generate_corpus(args.manifest, args.output):
        print(path)


if __name__ == "__main__":
    main()
